import requests
import urllib.parse
from concurrent.futures import ThreadPoolExecutor, as_completed
import logging
import json
import os
import uuid
import io
import datetime
from io import BytesIO
from dotenv import load_dotenv

# Load Environment Variables
load_dotenv()

from flask import Blueprint, render_template, request, jsonify, Response, stream_with_context, send_file, current_app
from flask_login import login_required, current_user
from firebase_admin import firestore
from pypdf import PdfReader
from bs4 import BeautifulSoup
from docx import Document

# Import module internal aplikasi
from app import firestore_db
from app.services.ai_service import AIService
from app.utils import ai_utils
from app.utils.citation_helper import generate_bibliography

# Inisialisasi Blueprint dan Logger
assistant_bp = Blueprint('assistant', __name__)
logger = logging.getLogger(__name__)

# ==============================================================================
# BAGIAN 1: HELPER LIMITATION CHECKER (RATE LIMITING)
# ==============================================================================

def check_limits(user, limit_type='generator'):
    """
    Mengecek apakah user sudah mencapai batas penggunaan harian.
    
    Args:
        user: Objek current_user Flask-Login
        limit_type: 'generator' (Max 3x) atau 'chat' (Max 4x)
    
    Returns:
        (Allowed: bool, Message: str)
    """
    # 1. BYPASS UNTUK PRO USER
    # Jika user punya flag is_pro = True, maka bebas limit
    if getattr(user, 'is_pro', False):
        return True, "Pro User - Unlimited Access"

    # 2. CEK DATABASE UNTUK FREE USER
    # Gunakan tanggal hari ini sebagai key agar reset setiap hari
    today_str = datetime.datetime.now().strftime('%Y-%m-%d')
    
    # Gunakan nama koleksi yang berbeda agar kuota tidak tercampur
    # usage_logs_gen  -> Kuota Tools Berat (Generator, Paraphrase, Outline)
    # usage_logs_chat -> Kuota Chat Ringan
    collection_name = 'usage_logs_gen' if limit_type == 'generator' else 'usage_logs_chat'
    max_limit = 3 if limit_type == 'generator' else 4
    
    # ID Dokumen: UserID_Tanggal
    doc_ref = firestore_db.collection(collection_name).document(f"{user.id}_{today_str}")
    doc = doc_ref.get()

    current_usage = 0
    if doc.exists:
        current_usage = doc.to_dict().get('count', 0)

    # Logika Pengecekan
    if current_usage >= max_limit:
        return False, f"Kuota Harian {limit_type.capitalize()} Habis ({current_usage}/{max_limit}). Upgrade ke Pro untuk akses tanpa batas!"
    
    return True, "OK"

def increment_limit(user, limit_type='generator'):
    """
    Menambah hitungan penggunaan (+1) setelah request sukses.
    Hanya dijalankan untuk User Free.
    """
    if getattr(user, 'is_pro', False): 
        return

    today_str = datetime.datetime.now().strftime('%Y-%m-%d')
    collection_name = 'usage_logs_gen' if limit_type == 'generator' else 'usage_logs_chat'
    
    doc_ref = firestore_db.collection(collection_name).document(f"{user.id}_{today_str}")
    
    if doc_ref.get().exists:
        # Jika dokumen ada, increment count
        doc_ref.update({'count': firestore.Increment(1)})
    else:
        # Jika dokumen belum ada (request pertama hari ini), buat baru
        doc_ref.set({
            'count': 1, 
            'userId': user.id, 
            'date': today_str,
            'email': getattr(user, 'email', 'unknown')
        })


# ==============================================================================
# BAGIAN 2: HALAMAN VIEW (FRONTEND RENDERING)
# ==============================================================================

@assistant_bp.route('/writing-assistant')
@login_required
def writing_assistant():
    """Halaman utama Writing Studio (React App Wrapper)."""
    project_id = request.args.get('id')
    project = None

    try:
        if project_id:
            doc_ref = firestore_db.collection('projects').document(project_id)
            doc = doc_ref.get()
            # Validasi kepemilikan project
            if doc.exists and doc.to_dict().get('userId') == str(current_user.id):
                project = doc.to_dict()
                project['id'] = doc.id
    except Exception as e:
        logger.error(f"Page Load Error: {e}")

    return render_template('writing_assistant.html', project=project)

@assistant_bp.route('/generator-kajian-teori')
@login_required
def generator_kajian_teori():
    """Halaman khusus workbench Kajian Teori (Bab 2)."""
    return render_template('generator_kajian_teori.html')

@assistant_bp.route('/thesis-defense')
@login_required
def thesis_defense_page():
    """Halaman simulasi sidang skripsi."""
    return render_template('thesis_defense.html')


# ==============================================================================
# BAGIAN 3: API MANAJEMEN PROJECT (CRUD)
# ==============================================================================

@assistant_bp.route('/api/projects/new', methods=['POST'])
@login_required
def api_create_project():
    """Membuat project baru."""
    try:
        new_project = {
            'userId': str(current_user.id),
            'title': 'Proyek Baru',
            'problem_statement': '',
            'methodology': '',
            'variables': '',
            'content': '', 
            'created_at': firestore.SERVER_TIMESTAMP,
            'updated_at': firestore.SERVER_TIMESTAMP
        }
        update_time, project_ref = firestore_db.collection('projects').add(new_project)
        return jsonify({'status': 'success', 'projectId': project_ref.id})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/projects', methods=['GET'])
@login_required
def get_user_projects():
    """Mengambil daftar project milik user."""
    try:
        docs = firestore_db.collection('projects')\
            .where('userId', '==', str(current_user.id))\
            .stream()
            
        projects = []
        for doc in docs:
            d = doc.to_dict()
            projects.append({
                'id': doc.id,
                'title': d.get('title', 'Tanpa Judul'),
                'updated_at': d.get('updated_at', '') 
            })
        
        # Sort client-side karena firestore composite index kadang ribet
        projects.sort(key=lambda x: str(x['updated_at']), reverse=True)
            
        return jsonify({'status': 'success', 'projects': projects})
    except Exception as e:
        logger.error(f"List Projects Error: {e}")
        return jsonify({'status': 'error', 'projects': [], 'message': str(e)}), 500

@assistant_bp.route('/api/projects/<project_id>', methods=['GET'])
@login_required
def get_project_details(project_id):
    """Mengambil detail satu project."""
    try:
        doc_ref = firestore_db.collection('projects').document(project_id)
        doc = doc_ref.get()
        
        if not doc.exists:
            return jsonify({'error': 'Project not found'}), 404
            
        data = doc.to_dict()
        if data.get('userId') != str(current_user.id):
            return jsonify({'error': 'Unauthorized'}), 403
            
        data['id'] = doc.id
        return jsonify({'status': 'success', 'project': data})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/projects/<project_id>', methods=['PUT'])
@login_required
def update_project(project_id):
    """Update data project (Auto-save)."""
    try:
        data = request.get_json()
        doc_ref = firestore_db.collection('projects').document(project_id)
        doc = doc_ref.get()
        
        if not doc.exists:
            return jsonify({'error': 'Project not found'}), 404
        if doc.to_dict().get('userId') != str(current_user.id):
            return jsonify({'error': 'Unauthorized'}), 403
            
        # Update timestamp
        data['updated_at'] = firestore.SERVER_TIMESTAMP
        doc_ref.update(data)
        
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==============================================================================
# BAGIAN 4: API REFERENCES & CITATIONS
# ==============================================================================

@assistant_bp.route('/api/project/<project_id>/references/add', methods=['POST'])
@login_required
def add_project_reference(project_id):
    """Menambah referensi ke project (Safe Mode dengan .add)."""
    try:
        ref_data = request.get_json()
        if not ref_data:
            return jsonify({'status': 'error', 'message': 'Data kosong'}), 400

        # Verifikasi kepemilikan project
        proj_ref = firestore_db.collection('projects').document(project_id)
        proj = proj_ref.get()
        if not proj.exists or proj.to_dict().get('userId') != str(current_user.id):
            return jsonify({'status': 'error', 'message': 'Unauthorized'}), 403

        # Inject Data Penting
        ref_data['projectId'] = project_id
        ref_data['userId'] = str(current_user.id)
        ref_data['createdAt'] = firestore.SERVER_TIMESTAMP
        
        # Bersihkan field ID jika ada (biar Firestore yang generate)
        if 'id' in ref_data:
            del ref_data['id']

        # Simpan ke Firestore
        update_time, doc_ref = firestore_db.collection('citations').add(ref_data)
        
        # Update dokumen agar punya field 'id' yang sama dengan doc ID (opsional tapi berguna untuk frontend)
        doc_ref.update({'id': doc_ref.id})

        return jsonify({'status': 'success', 'message': 'Referensi tersimpan', 'id': doc_ref.id}), 200

    except Exception as e:
        logger.error(f"Add Reference Error: {e}")
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/references/delete', methods=['POST'])
@login_required
def delete_reference():
    """Menghapus referensi."""
    try:
        data = request.get_json()
        ref_id = data.get('id')
        
        if not ref_id:
            return jsonify({'error': 'No Reference ID'}), 400
            
        doc_ref = firestore_db.collection('citations').document(ref_id)
        doc = doc_ref.get()
        
        if doc.exists and doc.to_dict().get('userId') == str(current_user.id):
            doc_ref.delete()
            return jsonify({'status': 'success'})
        else:
            return jsonify({'error': 'Not found or Unauthorized'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==============================================================================
# BAGIAN 5: API AI GENERATOR UTAMA (STREAMING)
# ==============================================================================

@assistant_bp.route('/api/writing-assistant', methods=['POST'])
@login_required
def api_writing_assistant():
    """
    ENDPOINT GENERATOR UTAMA (BAB, OUTLINE, PARAGRAF, DLL).
    Mendukung Streaming & Auto-Context Injection dari Project.
    """
    try:
        # 1. CEK LIMIT GENERATOR (Max 3x untuk Free User)
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed:
            return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        req_data = request.get_json()
        if not req_data:
            return jsonify({'error': 'No input data'}), 400

        # Persiapan Data Input
        input_payload = req_data.get('data', {})
        project_id = input_payload.get('projectId') or req_data.get('projectId')

        # 2. AUTO-LOAD CONTEXT DARI FIRESTORE
        # Jika ada project ID, kita ambil data skripsi user untuk dijadikan konteks AI
        if project_id:
            try:
                doc_ref = firestore_db.collection('projects').document(project_id)
                doc = doc_ref.get()
                
                if doc.exists:
                    p_data = doc.to_dict()
                    # Inject field ke payload AI
                    input_payload['context_title'] = p_data.get('title', '')
                    input_payload['context_problem'] = p_data.get('problem_statement', '')
                    input_payload['context_method'] = p_data.get('methodology', '')
                    input_payload['context_variables'] = p_data.get('variables_indicators', p_data.get('variables', ''))
                    input_payload['context_hypothesis'] = p_data.get('hypothesis', '')
                    input_payload['context_objectives'] = p_data.get('research_objectives', '')
                    input_payload['context_framework'] = p_data.get('theoretical_framework', '')
                    
                    # Update request data dengan payload yang sudah diperkaya
                    req_data['data'] = input_payload
                    logger.info(f"Context loaded for Project: {project_id}")
            except Exception as db_err:
                logger.warning(f"Context Auto-Load Failed (Non-Fatal): {db_err}")

        # 3. PANGGIL AI SERVICE (STREAMING)
        # Fungsi ini ada di ai_service.py dan akan return generator
        result = AIService.writing_assistant_stream(current_user, req_data)

        # 4. CATAT PEMAKAIAN (INCREMENT)
        # Kita catat penggunaan setelah berhasil memanggil service
        increment_limit(current_user, 'generator')

        # 5. HANDLE RETURN TYPE
        
        # KASUS A: Jika hasil berupa Dictionary (Biasanya dari tools non-stream)
        if isinstance(result, dict):
            return jsonify(result), 200
        
        # KASUS B: Jika hasil berupa String (Legacy Code)
        if isinstance(result, str):
            return jsonify({'generated_content': result}), 200

        # KASUS C: Jika hasil berupa Generator/Stream (Default Behavior)
        if hasattr(result, '__iter__'):
            return Response(stream_with_context(result), mimetype='text/html')

        # Fallback terakhir
        return jsonify({'generated_content': str(result)}), 200

    except Exception as e:
        logger.error(f"Writing API Error: {e}")
        return jsonify({'error': str(e)}), 500


# ==============================================================================
# BAGIAN 6: API CHAT (STREAMING)
# ==============================================================================

@assistant_bp.route('/chat/stream', methods=['POST'])
@login_required
def chat_with_ai_stream():
    """
    ENDPOINT CHAT KHUSUS (LIMIT 4x UNTUK FREE USER).
    Menggunakan limitasi kata max 200 (diatur di ai_service.py).
    """
    try:
        # 1. CEK LIMIT CHAT
        allowed, msg = check_limits(current_user, 'chat')
        if not allowed:
            return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        data = request.get_json()
        
        # 2. BUNGKUS DATA SESUAI FORMAT SERVICE
        # Kita set task='chat' secara eksplisit agar masuk logika prompt limit 200 kata di ai_service
        wrapper_data = {
            'task': 'chat',
            'data': {
                'message': data.get('message'),
                'context': data.get('context', ''), 
                'projectId': data.get('projectId'),
                'context_title': data.get('context_title', '') # Opsional
            }
        }
        
        # 3. PANGGIL SERVICE (UNIFIED)
        result = AIService.writing_assistant_stream(current_user, wrapper_data)
        
        # 4. INCREMENT LIMIT CHAT
        increment_limit(current_user, 'chat')

        return Response(stream_with_context(result), mimetype='text/plain')

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Endpoint Chat Alias (untuk kompatibilitas frontend lama)
@assistant_bp.route('/chat', methods=['POST'])
@login_required
def chat_with_ai(): 
    return chat_with_ai_stream()


# ==============================================================================
# BAGIAN 7: API TOOLS SPESIFIK (GENERATOR TOOLS)
# ==============================================================================

@assistant_bp.route('/api/generate-outline', methods=['POST'])
@login_required
def api_generate_outline():
    """Generate Outline Skripsi (Direct Call to Utils)."""
    try:
        # 1. Cek Limit
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed: 
            return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        # 2. Ambil Data
        data = request.get_json() or {}
        
        # 3. Construct Project Context Manual
        # Kita ambil judul dari payload frontend (judul_penelitian) atau default
        project_context = {
            'title': data.get('judul_penelitian', 'Topik Skripsi'),
            'problem_statement': 'Generate Outline Mode',
            'methodology': data.get('methodology', 'Umum')
        }
        
        # 4. Panggil AI Utils LANGSUNG (Bypass AIService)
        # Force task_type='generate_outline'
        stream_result = ai_utils.generate_academic_draft_stream(
            user=current_user,
            task_type='generate_outline',  # <--- KITA PAKSA DISINI
            input_data=data,
            project_context=project_context,
            selected_model='free_standard'
        )
        
        # 5. Consume Stream & Bersihkan Output
        full_text = ""
        for chunk in stream_result:
            if isinstance(chunk, bytes):
                full_text += chunk.decode('utf-8')
            else:
                full_text += str(chunk)
        
        # 6. Parse JSON dengan Aman
        try:
            # Helper ini akan membuang teks "Baik, berikut..." dan ambil JSON-nya saja
            clean_text = ai_utils.clean_json_output(full_text)
            outline_json = json.loads(clean_text)
            
            increment_limit(current_user, 'generator')
            return jsonify(outline_json)
            
        except json.JSONDecodeError:
            logger.error(f"Outline Parse Error. Raw: {full_text}")
            # Fallback darurat: jika gagal parse, kirim error biar frontend tau
            return jsonify({
                'error': 'Gagal format JSON',
                'raw': full_text
            }), 500

    except Exception as e:
        logger.error(f"Generate Outline Error: {e}")
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/paraphrase', methods=['POST'])
@login_required
def paraphrase():
    """API Streaming Paraphrase."""
    # Paraphrase masuk kuota generator
    allowed, msg = check_limits(current_user, 'generator')
    if not allowed:
        return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

    data = request.get_json()
    text = data.get('text')
    style = data.get('style', 'academic')
    
    if not text: return jsonify({'error': 'Text is empty'}), 400

    wrapper_data = {
        'task': 'paraphrase',
        'data': {
            'content': text,
            'style': style
        }
    }

    increment_limit(current_user, 'generator')

    result = AIService.writing_assistant_stream(current_user, wrapper_data)
    return Response(stream_with_context(result), mimetype='text/plain')

@assistant_bp.route('/expand-text', methods=['POST'])
@login_required
def expand_text_endpoint():
    """API Streaming Magic Expand."""
    allowed, msg = check_limits(current_user, 'generator')
    if not allowed: return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

    data = request.get_json()
    text = data.get('text')
    project_id = data.get('projectId')
    
    if not text: return jsonify({'error': 'Text is empty'}), 400

    # Optional: Ambil referensi dari project untuk memperkaya expand
    context_refs_str = ""
    if project_id:
        try:
            refs_query = firestore_db.collection('citations')\
                .where('projectId', '==', project_id)\
                .limit(5).stream()
            
            ref_list = []
            for doc in refs_query:
                d = doc.to_dict()
                ref_list.append(f"- {d.get('title')} ({d.get('author')}, {d.get('year')})")
            
            if ref_list:
                context_refs_str = "\n".join(ref_list)
        except Exception as e:
            logger.error(f"Context Ref Error: {e}")

    wrapper_data = {
        'task': 'expand_text',
        'data': {
            'content': text,
            'context': context_refs_str
        }
    }

    increment_limit(current_user, 'generator')

    result = AIService.writing_assistant_stream(current_user, wrapper_data)
    return Response(stream_with_context(result), mimetype='text/plain')

@assistant_bp.route('/api/ai/edit-text', methods=['POST'])
@login_required
def ai_edit_text():
    """Endpoint untuk Floating Toolbar (Shorten, Formalize, dll)."""
    try:
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed: return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        data = request.get_json()
        selected_text = data.get('text', '')
        mode = data.get('mode', 'paraphrase') # mode: shorten, formalize, etc.
        should_stream = data.get('stream', False)
        
        if not selected_text: return jsonify({'error': 'Text empty'}), 400
        
        # Mapping mode ke task type AI
        task_map = {
            'shorten': 'shorten_text',
            'formalize': 'formalize_text',
            'paraphrase': 'paraphrase',
            'fix_grammar': 'grammar_check'
        }
        task_type = task_map.get(mode, 'paraphrase')

        wrapper_data = {
            'task': task_type,
            'data': {'content': selected_text}
        }
        
        increment_limit(current_user, 'generator')
        
        result = AIService.writing_assistant_stream(current_user, wrapper_data)

        # Handle jika frontend minta JSON vs Stream
        if should_stream:
            return Response(stream_with_context(result), mimetype='text/plain')
        else:
            # Consume stream to string
            full_text = "".join([chunk for chunk in result])
            return jsonify({'status': 'success', 'result': full_text})

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ==============================================================================
# BAGIAN 8: API ANALYSIS & LOGIC (HEAVY TOOLS)
# ==============================================================================

@assistant_bp.route('/api/analyze-style', methods=['POST'])
@login_required
def api_analyze_style():
    """Menganalisis gaya penulisan dari file user."""
    if 'file' not in request.files: return jsonify({'error': 'No file'}), 400
    try:
        # Style analysis biasanya task sekali jalan, hitungannya generator
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed: return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403
        
        file = request.files['file']
        # Membaca file dan analisa (Logic ada di AIService)
        profile = AIService.analyze_style_from_file(current_user, file)
        
        increment_limit(current_user, 'generator')
        return jsonify({'status': 'success', 'style_profile': profile})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/logic-matrix', methods=['POST'])
@login_required
def logic_matrix_endpoint():
    """Membuat Matrix Konsistensi Logika (Masalah vs Kesimpulan)."""
    try:
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed: return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        data = request.get_json()
        problem = data.get('problem', '').strip()
        conclusion = data.get('conclusion', '').strip()
        
        if not problem or not conclusion:
            return jsonify({'status': 'error', 'message': 'Data kurang lengkap.'}), 400
            
        result = AIService.generate_logic_matrix(current_user, problem, conclusion)
        
        increment_limit(current_user, 'generator')
        return jsonify({'status': 'success', 'data': result})
        
    except Exception as e:
        logger.error(f"Logic Matrix Error: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@assistant_bp.route('/api/logic/check', methods=['POST'])
@login_required
def check_logic():
    """Logic Check Menyeluruh (Bab 1 vs Bab 5)."""
    try:
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed: return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        data = request.get_json()
        # AIService.check_logic_consistency biasanya return JSON
        result = AIService.check_logic_consistency(data)
        
        increment_limit(current_user, 'generator')
        return jsonify({'status': 'success', 'data': result})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@assistant_bp.route('/api/check-method-compliance', methods=['POST'])
@login_required
def check_method_compliance_endpoint():
    """Cek apakah teks sesuai dengan metodologi (Kuali/Kuanti)."""
    try:
        # Ini fitur analisis, kita bisa masukkan ke generator limit atau free
        # Mari masukkan ke generator limit
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed: return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        data = request.get_json()
        text = data.get('text', '').strip()
        method_mode = data.get('method_mode', 'quantitative') 
        
        if not text: return jsonify({'status': 'error'}), 400

        issues = AIService.check_method_compliance(text, method_mode)
        
        increment_limit(current_user, 'generator')
        return jsonify({'status': 'success', 'issues': issues, 'checked_method': method_mode})
        
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ==============================================================================
# BAGIAN 9: API DEFENSE & PPT (ADVANCED)
# ==============================================================================

@assistant_bp.route('/api/defense/<action>', methods=['POST'])
@login_required
def defense_endpoint(action):
    """
    Endpoint Simulasi Sidang.
    Action: 'start', 'answer', 'evaluate'.
    """
    try:
        # Simulasi sidang sangat berat, wajib cek limit
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed and action == 'start': # Cek limit hanya saat mulai sesi
             return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        data = request.get_json()
        result = AIService.thesis_defense_simulation(current_user, action, data)
        
        # Increment limit jika memulai sesi baru
        if action == 'start':
            increment_limit(current_user, 'generator')

        response_key = 'response' if action != 'evaluate' else 'report'
        return jsonify({'status': 'success', response_key: result})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/generate-ppt', methods=['POST'])
@login_required
def generate_ppt_endpoint():
    """Generate Powerpoint Slide."""
    try:
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed: return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        pptx_file = AIService.generate_ppt(current_user, request.get_json())
        filename = f"OnThesis_Slide_{datetime.datetime.now().strftime('%Y%m%d')}.pptx"
        
        increment_limit(current_user, 'generator')
        
        return send_file(
            pptx_file,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.presentationml.presentation'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==============================================================================
# BAGIAN 10: API UTILS (SEARCH, BIBLIOGRAPHY, EXPORT)
# ==============================================================================

@assistant_bp.route('/api/extract-pdf-simple', methods=['POST'])
@login_required
def extract_pdf_simple():
    """Ekstrak teks dari PDF upload."""
    try:
        if 'document' not in request.files:
            return jsonify({'status': 'error', 'error': 'No file'}), 400
            
        file = request.files['document']
        reader = PdfReader(file)
        full_text = []
        max_pages = min(len(reader.pages), 50) 
        
        for i in range(max_pages):
            page_text = reader.pages[i].extract_text()
            if page_text: full_text.append(page_text)
        
        joined_text = "\n".join(full_text)
        
        if len(joined_text.strip()) < 50:
             return jsonify({'status': 'error', 'error': 'Teks tidak terbaca (Scan?).'}), 422

        return jsonify({
            'status': 'success', 
            'data': {'filename': file.filename, 'content': joined_text, 'page_count': len(reader.pages)}
        })
    except Exception as e:
        return jsonify({'status': 'error', 'error': str(e)}), 500

@assistant_bp.route('/api/unified-search-references', methods=['POST'])
@login_required
def unified_search_references():
    """
    Search references dari Crossref/DOAJ/dll.
    Biasanya fitur ini gratis agar user mudah cari referensi.
    """
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        year_filter = str(data.get('year', '')).strip()
        
        if not (year_filter.isdigit() and len(year_filter) == 4): year_filter = None
        if not query: return jsonify({'status': 'error', 'message': 'Query kosong'}), 400

        HEADERS = {'User-Agent': 'Mozilla/5.0 (OnThesis Academic Bot)'}
        results = []

        # Logic search sederhana ke Crossref (Implementasi penuh bisa via Helper)
        params = {'query': query, 'rows': 5}
        if year_filter: params['filter'] = f'from-pub-date:{year_filter}'
        
        try:
            resp = requests.get("https://api.crossref.org/works", params=params, headers=HEADERS, timeout=5)
            if resp.status_code == 200:
                items = resp.json().get('message', {}).get('items', [])
                for item in items:
                    try:
                        title = item.get('title', [''])[0]
                        auths = item.get('author', [])
                        authors_str = ", ".join([f"{a.get('family', '')}" for a in auths[:2]])
                        pub_date = item.get('published-print', {}).get('date-parts', [[None]])[0][0]
                        
                        results.append({
                            'title': title, 
                            'author': authors_str, 
                            'source': 'Crossref', 
                            'year': str(pub_date) if pub_date else 'n.d.',
                            'doi': item.get('DOI', ''),
                            'url': item.get('URL', '')
                        })
                    except: continue
        except Exception as api_err:
            logger.warning(f"Search API Error: {api_err}")

        return jsonify({'status': 'success', 'results': results}) 

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@assistant_bp.route('/api/generate-bibliography', methods=['POST'])
@login_required
def api_generate_bibliography():
    """Generate Daftar Pustaka otomatis."""
    try:
        data = request.get_json()
        project_id = data.get('projectId')
        if not project_id: return jsonify({'error': 'No Project ID'}), 400
        
        refs_query = firestore_db.collection('citations').where('projectId', '==', project_id).stream()
        references = [doc.to_dict() for doc in refs_query]
        
        # Panggil helper function
        html_out, text_out = generate_bibliography(references)
        
        return jsonify({'status': 'success', 'html': html_out, 'text': text_out})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/export-docx', methods=['POST'])
@login_required
def export_docx_endpoint():
    """Export Project ke DOCX."""
    try:
        data = request.get_json()
        project_id = data.get('projectId')
        
        doc_ref = firestore_db.collection('projects').document(project_id)
        doc = doc_ref.get()
        if not doc.exists: return jsonify({'error': 'Not found'}), 404
        
        project_data = doc.to_dict()
        html_content = project_data.get('content', '')
        title = project_data.get('title', 'Draft Skripsi')

        # Setup Document
        document = Document()
        document.add_heading(title, 0)
        
        author_name = getattr(current_user, 'name', None) or getattr(current_user, 'username', 'Mahasiswa')
        document.add_paragraph(f"Penulis: {author_name}")
        document.add_paragraph(f"Generated by OnThesis AI")
        document.add_page_break()

        # Parse HTML content (Simple Version)
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Iterasi elemen penting saja
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'li']):
            text = element.get_text().strip()
            if not text: continue
            
            if element.name == 'h1': document.add_heading(text, level=1)
            elif element.name == 'h2': document.add_heading(text, level=2)
            elif element.name == 'h3': document.add_heading(text, level=3)
            elif element.name == 'li': document.add_paragraph(text, style='List Paragraph')
            else: document.add_paragraph(text)

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        
        safe_title = "".join([c for c in title if c.isalnum() or c in (' ', '_')]).rstrip()
        filename = f"{safe_title[:30]}_Draft.docx"

        return send_file(
            file_stream, 
            as_attachment=True, 
            download_name=filename, 
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/review-document', methods=['POST'])
def review_document():
    """
    Endpoint Review Document (AI Editor).
    Note: Biasanya diakses oleh plugin editor.
    """
    try:
        # Cek limit jika mau, tapi request ini dari plugin editor
        # Kita anggap generator usage
        if current_user.is_authenticated:
            allowed, msg = check_limits(current_user, 'generator')
            if not allowed: return jsonify({'status': 'error', 'message': msg}), 403
            increment_limit(current_user, 'generator')

        data = request.get_json() or {}
        text = data.get("text", "").strip()
        
        if not text:
            return jsonify({"error": "Tidak ada teks"}), 400

        # Panggil logika review (bisa via AIService atau langsung litellm)
        # Disini kita mock implementasi sederhana atau panggil service
        # Agar konsisten, kita panggil logic review dari AIService jika ada,
        # atau implementasi inline jika service belum siap.
        
        # Implementasi inline untuk review grammar/typo
        # (Idealnya dipindah ke AIService)
        from litellm import completion
        
        prompt = """
        You are a professional academic reviewer.
        TASK: Review the text for grammar errors, typo, and clarity issues.
        OUTPUT: JSON Array [{"target": "bad text", "issue": "explanation", "fix": "correction"}].
        NO MARKDOWN.
        """
        
        resp = completion(
            model="groq/llama-3.3-70b-versatile",
            messages=[{"role":"system", "content": prompt}, {"role":"user", "content": text}],
            response_format={"type": "json_object"}
        )
        
        content = resp.choices[0].message.content
        result_json = json.loads(ai_utils.clean_json_output(content))
        
        # Handle format beda dari LLM
        reviews = result_json if isinstance(result_json, list) else result_json.get('reviews', [])
        
        return jsonify({"status": "success", "reviews": reviews})

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@assistant_bp.route('/api/assistant/logic-check', methods=['POST'])
@login_required
def logic_check_route():
    """
    ENDPOINT: AUDIT KONSISTENSI (BENANG MERAH).
    Mengecek keselarasan antara Judul vs Masalah vs Tujuan.
    """
    try:
        # 1. CEK LIMIT (Fitur ini "Berat", masuk kuota generator)
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed:
            return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        data = request.get_json()
        
        # Validasi Input Minimal
        if not data.get('title') or not data.get('problem'):
            return jsonify({'error': 'Judul dan Masalah wajib diisi untuk audit.'}), 400

        # 2. PANGGIL AI SERVICE
        result = AIService.analyze_logic_flow(current_user, data)
        
        # 3. INCREMENT LIMIT (Jika sukses)
        increment_limit(current_user, 'generator')
        
        return jsonify(result)

    except Exception as e:
        logger.error(f"Logic Check Error: {e}")
        return jsonify({'error': str(e)}), 500

# ==============================================================================
# BAGIAN 10: API AI GENERATOR UTAMA (STREAMING)
# ==============================================================================
@assistant_bp.route('/api/assistant/generate-stream', methods=['POST'])
@login_required
def generate_stream_endpoint():
    """
    Endpoint khusus untuk Generator Tab baru (Orchestrator).
    Langsung menghubungkan Frontend ke ai_utils.generate_academic_draft_stream.
    """
    try:
        # 1. Cek Limit Generator
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed:
            return jsonify({'message': msg}), 403

        # 2. Ambil Data dari Frontend
        data = request.get_json()
        if not data:
            return jsonify({'message': 'No input data'}), 400

        task_type = data.get('task', 'general')
        
        # 3. Construct Project Context
        # Frontend baru mengirim key 'context_*' langsung di root payload
        project_context = {
            'title': data.get('context_title', ''),
            'problem_statement': data.get('context_problem', ''),
            'methodology': data.get('context_method', ''),
            'variables': data.get('context_variables', ''),
            # Tambahkan field lain jika perlu
        }

        # 4. Panggil AI Utils (Logic yang baru kita update)
        # Kita passing 'data' full sebagai 'input_data' agar parameter bab (method_mode, dll) terbaca
        result = ai_utils.generate_academic_draft_stream(
            user=current_user,
            task_type=task_type,
            input_data=data, 
            project_context=project_context,
            # Bisa hardcode model atau biarkan default di utils
            selected_model='free_standard', 
            editor_context=data.get('previous_content', '')
        )

        # 5. Catat Usage (Hanya jika stream berhasil diinisiasi)
        increment_limit(current_user, 'generator')

        return result

    except Exception as e:
        logger.error(f"Generate Stream Error: {e}")
        # Return JSON error dengan status 500
        return jsonify({'message': f"Server Error: {str(e)}"}), 500